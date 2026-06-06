import json
import os
import sys
from html.parser import HTMLParser
from urllib.parse import parse_qs, urlparse

from docx import Document
from docx.shared import Inches, RGBColor


class NoteHtmlParser(HTMLParser):
    def __init__(self, document, attachments_root):
        super().__init__()
        self.document = document
        self.attachments_root = attachments_root
        self.paragraph = None
        self.list_mode = None
        self.bold = False
        self.italic = False
        self.color = None
        self.highlight = False

    def ensure_paragraph(self, style=None):
        if self.paragraph is None:
            self.paragraph = self.document.add_paragraph(style=style)
        return self.paragraph

    def close_paragraph(self):
        self.paragraph = None

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        if tag in ("p", "div"):
            self.ensure_paragraph()
        elif tag == "h2":
            self.ensure_paragraph("Heading 2")
        elif tag == "h3":
            self.ensure_paragraph("Heading 3")
        elif tag == "br":
            self.ensure_paragraph().add_run().add_break()
        elif tag in ("strong", "b"):
            self.bold = True
        elif tag in ("em", "i"):
            self.italic = True
        elif tag == "span":
            style = attrs_dict.get("style", "")
            self.color = extract_css_color(style, "color") or self.color
            self.highlight = bool(extract_css_color(style, "background-color") or extract_css_color(style, "background"))
        elif tag == "font":
            self.color = attrs_dict.get("color") or self.color
        elif tag == "ul":
            self.list_mode = "List Bullet"
        elif tag == "ol":
            self.list_mode = "List Number"
        elif tag == "li":
            self.ensure_paragraph(self.list_mode or "List Bullet")
        elif tag == "img":
            self.insert_image(attrs_dict.get("src", ""))

    def handle_endtag(self, tag):
        if tag in ("p", "div", "h2", "h3", "li"):
            self.close_paragraph()
        elif tag in ("strong", "b"):
            self.bold = False
        elif tag in ("em", "i"):
            self.italic = False
        elif tag == "span":
            self.color = None
            self.highlight = False
        elif tag == "font":
            self.color = None
        elif tag in ("ul", "ol"):
            self.list_mode = None

    def handle_data(self, data):
        if not data:
            return
        run = self.ensure_paragraph(self.list_mode).add_run(data)
        run.bold = self.bold
        run.italic = self.italic
        if self.color:
            rgb = parse_hex_color(self.color)
            if rgb:
                run.font.color.rgb = RGBColor(*rgb)
        if self.highlight:
            # python-docx supports named highlight colors only; yellow is the practical default.
            run.font.highlight_color = 7

    def insert_image(self, src):
        image_path = image_path_from_src(src, self.attachments_root)
        if not image_path or not os.path.exists(image_path):
            return
        paragraph = self.ensure_paragraph()
        run = paragraph.add_run()
        try:
            run.add_picture(image_path, width=Inches(5.8))
        except Exception:
            pass
        self.close_paragraph()


def extract_css_color(style, name):
    for chunk in style.split(";"):
        if ":" not in chunk:
            continue
        key, value = chunk.split(":", 1)
        if key.strip().lower() == name:
            return value.strip()
    return None


def parse_hex_color(value):
    value = value.strip()
    if not value.startswith("#"):
        return None
    value = value[1:]
    if len(value) == 3:
        value = "".join(ch * 2 for ch in value)
    if len(value) != 6:
        return None
    try:
        return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))
    except ValueError:
        return None


def image_path_from_src(src, attachments_root):
    parsed = urlparse(src)
    if parsed.path != "/api/attachment":
        return None
    query = parse_qs(parsed.query)
    lesson_id = (query.get("lessonId") or [""])[0]
    file_name = (query.get("file") or [""])[0]
    if not lesson_id or not file_name:
        return None
    safe_lesson = "".join(ch if ch.isalnum() or ch in "-_" else "-" for ch in lesson_id)
    safe_file = os.path.basename(file_name)
    return os.path.join(attachments_root, safe_lesson, safe_file)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: export_note_docx.py payload.json output.docx")

    with open(sys.argv[1], "r", encoding="utf-8") as handle:
        payload = json.load(handle)

    document = Document()
    title = payload.get("title") or "Lesson Notes"
    document.add_heading(title, level=1)

    parser = NoteHtmlParser(document, payload["attachmentsRoot"])
    parser.feed(payload.get("content") or "")

    document.save(sys.argv[2])


if __name__ == "__main__":
    main()
